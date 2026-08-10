"""Export results to CSV, Excel, or Word. Datasets -> CSV/Excel; reports -> Word."""
from __future__ import annotations

from pathlib import Path

import pandas as pd


def _export_safe(df: pd.DataFrame) -> pd.DataFrame:
    """Make text safe for the next tool (Excel, Power BI, other CSV readers):
    remove carriage returns and newlines inside fields (a common cause of
    'EOF within quoted string'), drop the Unicode replacement char and other
    control characters. Meaning is preserved; only breakage is removed."""
    import re as _r
    out = df.copy()
    ctrl = _r.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\ufffd]")
    for c in out.columns:
        out[c] = out[c].map(lambda v: v if not isinstance(v, str)
                            else ctrl.sub("", v.replace("\r", "").replace("\n", " ")))
    return out


def to_csv(df: pd.DataFrame, path: Path) -> Path:
    _export_safe(df).to_csv(path, index=False)
    return path


def to_xlsx(df: pd.DataFrame, path: Path, sheet: str = "Result") -> Path:
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name=sheet[:31], index=False)
    return path


def to_docx(title: str, df: pd.DataFrame, path: Path, intro: str = "") -> Path:
    """A clean Word report: title, optional intro, then the data as a table."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)
    if intro:
        doc.add_paragraph(intro)
    if df is None or df.empty:
        doc.add_paragraph("No records to report.")
        doc.save(path)
        return path
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for j, c in enumerate(cols):
        cell = table.rows[0].cells[j]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for _, row in df.head(2000).iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            cells[j].text = "" if pd.isna(row[c]) else str(row[c])
    if len(df) > 2000:
        doc.add_paragraph(f"… and {len(df) - 2000} more rows (full data available as CSV/Excel).")
    doc.save(path)
    return path


def export(df: pd.DataFrame, fmt: str, path: Path, title: str = "Report", intro: str = "") -> Path:
    fmt = fmt.lower()
    if fmt == "csv":
        return to_csv(df, path)
    if fmt in ("xlsx", "excel"):
        return to_xlsx(df, path)
    if fmt in ("docx", "word"):
        return to_docx(title, df, path, intro)
    raise ValueError(f"Unknown format: {fmt}")
